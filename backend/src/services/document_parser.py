"""
Document parsing and text extraction service.
Supports multiple formats: PDF, DOCX, HTML, TXT.
"""

import os
import re
import mimetypes
from pathlib import Path
from typing import Tuple, Dict, Any, Optional
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles document parsing for multiple formats."""

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.html', '.htm', '.txt'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file format is supported."""
        _, ext = os.path.splitext(filename.lower())
        return ext in DocumentParser.SUPPORTED_FORMATS

    @staticmethod
    def parse(file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document and extract text and metadata.
        
        Args:
            file_path: Path to the document file
            file_type: File extension (pdf, docx, html, txt)
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        file_type = file_type.lower().strip('.')
        
        if file_type == 'pdf':
            return DocumentParser._parse_pdf(file_path)
        elif file_type == 'docx':
            return DocumentParser._parse_docx(file_path)
        elif file_type in ['html', 'htm']:
            return DocumentParser._parse_html(file_path)
        elif file_type == 'txt':
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _parse_txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            
            metadata = {
                'format': 'text',
                'pages': 1,
                'language': 'en',
                'word_count': len(content.split()),
            }
            
            return content, metadata
        except Exception as e:
            logger.error(f"Error parsing TXT file {file_path}: {str(e)}")
            raise

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF file using PyPDF2."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 required for PDF parsing. Install: pip install PyPDF2")
        
        try:
            content_parts = []
            metadata = {'format': 'pdf', 'pages': 0}
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                metadata['pages'] = num_pages
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        content_parts.append(f"\n--- Page {page_num} ---\n{text}")
                
                if not content_parts:
                    logger.warning(f"No text extracted from PDF {file_path}. It might be scanned.")
                    content_parts.append("[SYSTEM WARNING: No text could be extracted from this document. It may be a scanned PDF or image-based. Text extraction requires searchable text.]")
                
                # Extract document info
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['subject'] = pdf_reader.metadata.get('/Subject', '')
            
            content = '\n'.join(content_parts)
            metadata['word_count'] = len(content.split())
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
            raise

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX file using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for DOCX parsing. Install: pip install python-docx")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n'.join(paragraphs)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_text.append(' | '.join(row_cells))
                content += '\n\n' + '\n'.join(table_text)
            
            # Extract metadata
            metadata = {
                'format': 'docx',
                'core_properties': {},
            }
            
            if doc.core_properties:
                props = doc.core_properties
                metadata['core_properties'] = {
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': str(props.created) if props.created else None,
                }
            
            metadata['word_count'] = len(content.split())
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            raise

    @staticmethod
    def _parse_html(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("BeautifulSoup4 required for HTML parsing. Install: pip install beautifulsoup4")
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Extract metadata
            metadata = {'format': 'html'}
            
            title_tag = soup.find('title')
            if title_tag:
                metadata['title'] = title_tag.get_text(strip=True)
            
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                if tag.get('name') == 'author':
                    metadata['author'] = tag.get('content', '')
                elif tag.get('name') == 'description':
                    metadata['description'] = tag.get('content', '')
            
            metadata['word_count'] = len(content.split())
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error parsing HTML file {file_path}: {str(e)}")
            raise


class DocumentChunker:
    """Chunks documents for indexing and retrieval."""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 100):
        """
        Initialize chunker.
        
        Args:
            chunk_size: Number of words per chunk
            overlap: Number of overlapping words between chunks
        """
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> list:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Full document text
            metadata: Document metadata
            
        Returns:
            List of chunk dictionaries
        """
        # Split by sentences first
        sentences = self._split_sentences(text)
        
        chunks = []
        current_chunk = []
        current_word_count = 0
        page_number = 1
        
        for i, sentence in enumerate(sentences):
            words_in_sentence = len(sentence.split())
            
            if current_word_count + words_in_sentence > self.chunk_size and current_chunk:
                # Save current chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'page_number': self._estimate_page(i, len(sentences)),
                    'section': self._detect_section(chunk_text),
                    'word_count': len(chunk_text.split()),
                })
                
                # Create overlap
                overlap_sentences = []
                word_count = 0
                for j in range(len(current_chunk) - 1, -1, -1):
                    overlap_sentences.insert(0, current_chunk[j])
                    word_count += len(current_chunk[j].split())
                    if word_count > self.overlap:
                        break
                
                current_chunk = overlap_sentences
                current_word_count = word_count
            
            current_chunk.append(sentence)
            current_word_count += words_in_sentence
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'text': chunk_text,
                'page_number': page_number,
                'section': self._detect_section(chunk_text),
                'word_count': len(chunk_text.split()),
            })
        
        return chunks

    @staticmethod
    def _split_sentences(text: str) -> list:
        """Split text into sentences."""
        # Simple sentence splitter
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]

    @staticmethod
    def _estimate_page(sentence_index: int, total_sentences: int, avg_page_size: int = 250) -> int:
        """Estimate page number based on sentence position."""
        return max(1, (sentence_index * 100) // max(1, total_sentences))

    @staticmethod
    def _detect_section(text: str, max_length: int = 50) -> str:
        """Detect section title from chunk start."""
        lines = text.split('\n')
        for line in lines[:3]:
            line = line.strip()
            if line and len(line) < max_length and line.isupper():
                return line
            if line and ':' in line:
                return line[:max_length]
        return "Main"
